import os
import random
import re
import time

import requests
from bs4 import BeautifulSoup
import docx


def downpic(url,header):

    #设置图片的保存地址
    path = 'F:/work/python/pyqt5/imgs'
    if not os.path.exists(path):
        os.mkdir(path)
    #下载图片并保存到本地
    imgname = url.split('/')[-1]
    filename = os.path.join(path, imgname)
    if url.endswith('.png') or url.endswith('jpg') or url.endswith('gif'):
        r = requests.get(url,headers=header)
        try:
            fw = open(filename, 'wb')  # 指定绝对路径
            fw.write(r.content)  # 保存图片到本地指定目录
        except Exception as e:
            print("Error: "+str(e))

    return filename

def get():
    urls = ["http://www.xdbcb8.com/pyqt5/"]
    for i in range(2, 11):
        urls.append("http://www.xdbcb8.com/pyqt5/page/"+str(i))
    for url in urls:
        text = requests.get(url)
        soup = BeautifulSoup(text.text, 'html.parser')
        #设置消息头的User-Agent
        my_headers = [
            "Mozilla/5.0 (Windows NT 6.3; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/39.0.2171.95 Safari/537.36",
            "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_9_2) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/35.0.1916.153 Safari/537.36",
            "Mozilla/5.0 (Windows NT 6.1; WOW64; rv:30.0) Gecko/20100101 Firefox/30.0"
            "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_9_2) AppleWebKit/537.75.14 (KHTML, like Gecko) Version/7.0.3 Safari/537.75.14",
            "Mozilla/5.0 (compatible; MSIE 10.0; Windows NT 6.2; Win64; x64; Trident/6.0)"
            ]

        #获取每节课的内容
        for item in soup.find_all("a", text="阅读全文"):
        # if True:
            doc = docx.Document()
            #获取一节课的url
            #newurl = item.get('href')
            newurl = 'http://www.xdbcb8.com/archives/98.html'
            #设置消息头
            header = {
                "Accept": "*/*",
                "Accept-Encoding": "gzip,deflate",
                "Connection": "keep-alive",
                "User-Agent": random.choice(my_headers),
                "Referer":  newurl
            }
            #获取一节课的内容
            itemtext = requests.get(newurl)

            newsoup = BeautifulSoup(itemtext.text,'html.parser')
            #获取这节课的课名
            title = newsoup.find_all('h1',class_ = 'entry-title')
            doc.add_heading(title[0].text,0)

            for newitem in newsoup.find_all(['h4','h5','p','pre','img']):
                begin = [" 学点编程吧", "搜索热点", "文章来源：www.xdbcb8.com，转载请注明出处。"]
                end = ['最后','结束', "ok，今天就到这里，我们下期再约。","点赞"]
                #过滤开头和结尾不要的内容
                if newitem.text in begin:
                    continue
                elif newitem.string not in end :
                    if newitem.name == 'h4':
                        doc.add_paragraph().add_run('    ' + newitem.text).bold = True
                    elif newitem.name == 'h5':
                        doc.add_paragraph().add_run(newitem.text).bold = True
                    elif newitem.name == 'pre':
                        doc.add_paragraph(newitem.text)
                    elif newitem.name == 'p':
                        doc.add_paragraph(newitem.text)
                    elif newitem.name == 'img':
                        #获取图片的url
                        if newitem.attrs.get('class') != None:
                                picurl = newitem.attrs.get('data-original')
                                filename = downpic(picurl,header)
                                doc.add_picture(filename, width=docx.shared.Cm(17),height=docx.shared.Cm(8))

                    else:
                        pass
                else:
                    break
            try:
                doc.save('f:/work/python/pyqt5/' + title[0].text + '.docx')
            except Exception as e:
                print("Error: " + str(e))

            time.sleep(random.randint(1,5))


# get()


def kk():
    with open("sduview.html",'rb') as file :
        newsoup = BeautifulSoup(file, 'html.parser')
        title = newsoup.find_all('h1',class_ = 'entry-title')

        for item in newsoup.find_all(['h4','h5','p','pre','img']):
            if item.text in [" 学点编程吧","搜索热点","文章来源：www.xdbcb8.com，转载请注明出处。"]:
                continue
            elif item.string != '最后':
                if item.name == 'h4':
                    print(item.text)
                elif item.name == 'h5':
                    print(item.text)
                elif item.name == 'p':
                    print(item.text)
                elif item.name == 'pre':
                    print(item.text)
                elif item.name == 'img':
                    if item.attrs.get('class') != None:

                        print(item.attrs.get('data-original'))
                else:
                    pass
            else:
                break
kk()

